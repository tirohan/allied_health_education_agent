from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from backend.app.services.educator import ROLE_LABELS, EducatorRole, build_curriculum_outline


def build_teaching_pack(
    graph: dict[str, Any],
    query: str,
    role: str | None = None,
) -> dict[str, Any]:
    nodes = graph.get("nodes", [])
    by_type: dict[str, list[dict[str, Any]]] = {}
    for node in nodes:
        by_type.setdefault(str(node.get("entity_type")), []).append(node)

    resources = by_type.get("Resource", [])[:5]
    papers = by_type.get("Paper", [])[:2]
    cases = by_type.get("SimulationCase", [])[:1]
    counties = by_type.get("County", [])[:3]
    outline = build_curriculum_outline(graph, query, role)

    citations = []
    for node in resources + papers + cases + counties:
        citations.append(
            {
                "label": node.get("label"),
                "type": node.get("entity_type"),
                "source": f"{node.get('source_table')}.{node.get('source_id')}",
                "url": node.get("source_url"),
                "evidence": node.get("verification_status"),
            }
        )

    return {
        "title": "Allied Health Teaching Pack",
        "generated_at": datetime.now(UTC).isoformat(),
        "role": ROLE_LABELS.get(EducatorRole(role), "Educator") if role else "Educator",
        "planning_question": query,
        "resources": [item.get("label") for item in resources],
        "papers": [item.get("label") for item in papers],
        "simulation_case": cases[0].get("label") if cases else None,
        "county_indicators": [item.get("label") for item in counties],
        "citations": citations,
        "curriculum_outline": outline,
        "how_to_use": [
            "Share the pack with your teaching team before course planning.",
            "Pick one reading and one resource for pre-class preparation.",
            "Use the simulation case for skills practice or IPE discussion.",
            "Bring county context into debriefing and reflection prompts.",
        ],
    }


def teaching_pack_to_markdown(pack: dict[str, Any]) -> str:
    lines = [
        f"# {pack['title']}",
        "",
        f"Role: {pack.get('role')}",
        f"Planning question: {pack.get('planning_question')}",
        f"Generated: {pack.get('generated_at')}",
        "",
        "## Teaching resources",
    ]
    for item in pack.get("resources", []):
        lines.append(f"- {item}")
    lines.extend(["", "## Research readings"])
    for item in pack.get("papers", []):
        lines.append(f"- {item}")
    lines.extend(
        [
            "",
            "## Simulation case",
            f"- {pack.get('simulation_case') or 'No simulation case selected'}",
            "",
            "## Community context",
        ]
    )
    for item in pack.get("county_indicators", []):
        lines.append(f"- {item}")
    outline = pack.get("curriculum_outline") or {}
    lines.extend(["", "## Learning objectives"])
    for item in outline.get("learning_objectives", []):
        lines.append(f"- {item}")
    lines.extend(["", "## Suggested sequence"])
    for item in outline.get("suggested_sequence", []):
        lines.append(f"- {item}")
    lines.extend(["", "## Citations"])
    for citation in pack.get("citations", []):
        lines.append(
            f"- {citation.get('label')} ({citation.get('type')}, {citation.get('evidence')})"
        )
    lines.extend(["", "## How to use"])
    for item in pack.get("how_to_use", []):
        lines.append(f"- {item}")
    return "\n".join(lines) + "\n"


def teaching_pack_to_docx_bytes(pack: dict[str, Any]) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for Word export") from exc

    document = Document()
    document.add_heading(pack["title"], level=1)
    document.add_paragraph(f"Role: {pack.get('role')}")
    document.add_paragraph(f"Planning question: {pack.get('planning_question')}")
    document.add_paragraph(f"Generated: {pack.get('generated_at')}")

    document.add_heading("Teaching resources", level=2)
    for item in pack.get("resources", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Research readings", level=2)
    for item in pack.get("papers", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Simulation case", level=2)
    document.add_paragraph(str(pack.get("simulation_case") or "No simulation case selected"))

    document.add_heading("Community context", level=2)
    for item in pack.get("county_indicators", []):
        document.add_paragraph(str(item), style="List Bullet")

    outline = pack.get("curriculum_outline") or {}
    document.add_heading("Learning objectives", level=2)
    for item in outline.get("learning_objectives", []):
        document.add_paragraph(str(item), style="List Bullet")

    document.add_heading("Citations", level=2)
    for citation in pack.get("citations", []):
        document.add_paragraph(
            f"{citation.get('label')} ({citation.get('type')}, {citation.get('evidence')})",
            style="List Bullet",
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
